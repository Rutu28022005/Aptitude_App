"""
Updated script to create Word document with individual mobile wireframes and shorter conclusion
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Create a new Word document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('LAB 2 – Wireframing & App UI Flow Planning', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Aptitude App for Placement Preparation', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Introduction",
    "2. App Overview",
    "3. Screen Identification",
    "4. Low-Fidelity Wireframes",
    "5. UI Flow Diagram",
    "6. Navigation Logic",
    "7. Conclusion"
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Introduction
doc.add_heading('Introduction', level=1)
intro_text = """This document presents the wireframing and UI flow planning for the Aptitude App, a comprehensive mobile application designed for placement preparation. The wireframes demonstrate the visual structure, layout organization, and component placement for all essential screens without focusing on final styling or colors.

The UI flow diagram illustrates the complete navigation pattern, ensuring smooth and logical user experience throughout the application."""
doc.add_paragraph(intro_text)

doc.add_page_break()

# App Overview
doc.add_heading('App Overview', level=1)
doc.add_paragraph('Application Name: Aptitude App', style='List Bullet')
doc.add_paragraph('Purpose: A Flutter-based mobile application for aptitude test preparation targeting placement exams', style='List Bullet')

doc.add_paragraph('Core Features:', style='Heading 3')
features = [
    'AI-generated aptitude quizzes',
    'Multiple subjects: Mathematics/Quants, Logical Reasoning, VARC',
    'Progress tracking and performance analytics',
    'Adaptive difficulty levels',
    'Comprehensive quiz history'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# Screen Identification
doc.add_heading('Screen Identification', level=1)
doc.add_paragraph('Based on the app requirements and use cases, the following screens have been designed:')

# Create table
table = doc.add_table(rows=11, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = '#'
header_cells[1].text = 'Screen Name'
header_cells[2].text = 'Purpose'

# Data rows
screens_data = [
    ('1', 'Splash Screen', 'App branding and initial loading'),
    ('2', 'Login Screen', 'User authentication'),
    ('3', 'Sign Up Screen', 'New user registration'),
    ('4', 'Home/Dashboard Screen', 'Main navigation hub and subject selection'),
    ('5', 'Quiz Configuration Screen', 'Quiz setup and customization'),
    ('6', 'Quiz Screen', 'Active quiz-taking interface'),
    ('7', 'Result Screen', 'Quiz performance display'),
    ('8', 'History Screen', 'Past quiz records and review'),
    ('9', 'Analytics/Performance Screen', 'Performance metrics and insights'),
    ('10', 'Profile Screen', 'User account and settings')
]

for i, (num, name, purpose) in enumerate(screens_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = num
    row_cells[1].text = name
    row_cells[2].text = purpose

doc.add_paragraph()
doc.add_paragraph('Total Screens: 10 ✓ (Exceeds minimum requirement of 6-8 screens)', style='Intense Quote')

doc.add_page_break()

# Low-Fidelity Wireframes
doc.add_heading('Low-Fidelity Wireframes', level=1)

# Path to images
brain_dir = r'C:\Users\Admin\.gemini\antigravity\brain\4dcba3c2-fab2-42fe-891d-08bb8db95ba7'

# Individual wireframes with descriptions
wireframes = [
    {
        'title': '1. Splash Screen',
        'image': 'wireframe_splash_screen_1770011579889.png',
        'components': [
            'App logo placeholder (centered)',
            'App name text',
            'Loading indicator',
            'Layout: Centered vertical alignment with minimal elements'
        ]
    },
    {
        'title': '2. Login Screen',
        'image': 'wireframe_login_screen_1770011598447.png',
        'components': [
            'AppBar with "Login" title',
            'Email input field',
            'Password input field',
            'Login button',
            '"Forgot Password?" link',
            '"Don\'t have account? Sign Up" link',
            'Layout: Vertical form layout with clear call-to-action'
        ]
    },
    {
        'title': '3. Sign Up Screen',
        'image': 'wireframe_signup_screen_1770011618358.png',
        'components': [
            'AppBar with "Sign Up" title',
            'Name input field',
            'Email input field',
            'Password input field',
            'Sign Up button',
            '"Already have account? Login" link',
            'Layout: Extended vertical form with registration fields'
        ]
    },
    {
        'title': '4. Home/Dashboard Screen',
        'image': 'wireframe_home_screen_1770011637651.png',
        'components': [
            'AppBar with app title and profile icon',
            'Welcome message section',
            'Three subject cards (Mathematics, Logical Reasoning, VARC)',
            'Each card contains: Icon placeholder, Subject title, "Start Quiz" button',
            'Bottom navigation bar (Home, History, Analytics, Profile)',
            'Layout: Card-based grid layout with bottom navigation'
        ]
    },
    {
        'title': '5. Quiz Configuration Screen',
        'image': 'wireframe_quiz_config_1770011656991.png',
        'components': [
            'AppBar with "Quiz Setup" title and back button',
            'Subject selector dropdown',
            'Difficulty selector dropdown (Easy/Medium/Hard)',
            'Number of questions input',
            'Timer setting slider',
            '"Start Quiz" button',
            'Layout: Vertical form layout with clear hierarchy'
        ]
    },
    {
        'title': '6. Quiz Screen',
        'image': 'wireframe_quiz_screen_1770011676012.png',
        'components': [
            'AppBar with timer display and question counter',
            'Question card with placeholder text',
            'Four option buttons (A, B, C, D)',
            'Progress bar',
            '"Skip" and "Next" buttons',
            'Layout: Single question focus with clear option selection'
        ]
    },
    {
        'title': '7. Result Screen',
        'image': 'wireframe_result_screen_1770011696557.png',
        'components': [
            'AppBar with "Quiz Results" title',
            'Circular score display (percentage)',
            'Stats section (Correct, Wrong, Skipped, Time Taken)',
            'Question review expandable list',
            '"Review Answers" and "Back to Home" buttons',
            'Layout: Summary-focused with detailed breakdown'
        ]
    },
    {
        'title': '8. History Screen',
        'image': 'wireframe_history_screen_1770011713709.png',
        'components': [
            'AppBar with "Quiz History" title',
            'Date range filter section',
            'Scrollable quiz history cards',
            'Each card shows: Subject, Score percentage, Date, "View Details" button',
            'Bottom navigation bar (History active)',
            'Layout: List-based layout with filtering options'
        ]
    },
    {
        'title': '9. Analytics/Performance Screen',
        'image': 'wireframe_analytics_screen_1770011733198.png',
        'components': [
            'AppBar with "Performance Analytics" title',
            'Performance chart placeholder',
            'Stats grid (Total Quizzes, Average Score, Best Subject, Improvement Rate)',
            'Subject-wise breakdown list',
            'Bottom navigation bar (Analytics active)',
            'Layout: Data visualization focused with key metrics'
        ]
    },
    {
        'title': '10. Profile Screen',
        'image': 'wireframe_profile_screen_1770011752600.png',
        'components': [
            'AppBar with "Profile" title',
            'Circular profile photo placeholder',
            'User name and email',
            'Personal stats (Quizzes Taken, Average Score, Streak Days)',
            'Settings menu items (Account Settings, Notifications, Privacy, About)',
            'Logout button',
            'Bottom navigation bar (Profile active)',
            'Layout: User-centric with organized settings menu'
        ]
    }
]

# Add each wireframe
for wf in wireframes:
    doc.add_heading(wf['title'], level=2)
    
    image_path = os.path.join(brain_dir, wf['image'])
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(3.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Components:', style='Heading 4')
    for comp in wf['components']:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_paragraph()  # Add spacing

doc.add_page_break()

# UI Flow Diagram
doc.add_heading('UI Flow Diagram', level=1)

image5 = os.path.join(brain_dir, 'ui_flow_diagram_1770000422380.png')
if os.path.exists(image5):
    doc.add_picture(image5, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Flow Description', level=2)

doc.add_paragraph('Entry Point:', style='Heading 3')
doc.add_paragraph('App starts at Splash Screen', style='List Bullet')
doc.add_paragraph('Auto-navigates to Login Screen after 2 seconds', style='List Bullet')

doc.add_paragraph('Authentication Flow:', style='Heading 3')
doc.add_paragraph('Users can login or navigate to Sign Up Screen', style='List Bullet')
doc.add_paragraph('After successful registration, users return to Login', style='List Bullet')
doc.add_paragraph('Successful login directs to Home Screen', style='List Bullet')

doc.add_paragraph('Main Navigation Hub (Home Screen):', style='Heading 3')
doc.add_paragraph('Central hub with access to all major features', style='List Bullet')
doc.add_paragraph('Subject cards lead to Quiz Configuration Screen', style='List Bullet')
doc.add_paragraph('Bottom navigation provides access to History, Analytics, and Profile screens', style='List Bullet')

doc.add_paragraph('Quiz Flow (Primary User Journey):', style='Heading 3')
doc.add_paragraph('1. Select subject from Home → Quiz Config Screen', style='List Number')
doc.add_paragraph('2. Configure quiz settings → Quiz Screen', style='List Number')
doc.add_paragraph('3. Complete quiz → Result Screen', style='List Number')
doc.add_paragraph('4. Option to review answers (loops back to Quiz Screen) or return to Home', style='List Number')

doc.add_page_break()

# Navigation Logic
doc.add_heading('Navigation Logic', level=1)

nav_text = """The Aptitude App uses a hub-and-spoke navigation pattern with the Home screen as the central hub. Users authenticate via Splash and Login screens, then access three main areas through bottom navigation: quiz features, performance tracking, and profile management.

The quiz journey follows a linear path (Home → Config → Quiz → Results) with branching options to review answers or return home. Persistent bottom navigation enables quick switching between History, Analytics, and Profile sections. All screens provide intuitive back buttons and action buttons for context-specific navigation.

Key navigation features include minimal clicks to core actions (only 2 clicks from home to start a quiz), clear screen hierarchy, multiple access paths to features, and contextual next steps on each screen."""

doc.add_paragraph(nav_text)

doc.add_page_break()

# Conclusion (Shortened)
doc.add_heading('Conclusion', level=1)

conclusion_text = """This wireframing exercise established a comprehensive visual blueprint for the Aptitude App covering 10 screens across authentication (3 screens), core quiz features (4 screens), progress tracking (2 screens), and user management (1 screen).

The wireframes follow low-fidelity design principles with clear component hierarchy, logical organization, and user-friendly layouts. The UI flow diagram demonstrates smooth navigation, logical user journeys, and clear entry/exit points.

These wireframes will serve as the foundation for implementing UI components in Flutter, building navigation structure, and creating a consistent user experience."""

doc.add_paragraph(conclusion_text)

doc.add_paragraph()
doc.add_paragraph('Documentation Status: Complete and ready for submission ✓', style='Intense Quote')

doc.add_page_break()

# Document metadata
doc.add_paragraph('Prepared by: [Your Name]')
doc.add_paragraph('Roll Number: [Your Roll Number]')
doc.add_paragraph('Date: February 2, 2026')
doc.add_paragraph('Lab: LAB 2 – Wireframing & App UI Flow Planning')

# Save the document
output_path = r'C:\Users\Admin\Desktop\APP\Lab2_Wireframes_UIFlow_Updated.docx'
doc.save(output_path)

print(f"Updated Word document created successfully at: {output_path}")
